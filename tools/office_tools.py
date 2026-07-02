"""
Tools untuk membuat dan membaca dokumen Microsoft Office (Word & Excel).
Tidak butuh Word/Excel terinstall - langsung manipulasi file lewat library Python.

Semua path tetap melalui resolve_safe_path() dari file_tools untuk sandbox security.
"""

from tools.file_tools import resolve_safe_path, SandboxViolationError

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from openpyxl import Workbook, load_workbook
except ImportError:
    Workbook = None
    load_workbook = None


def create_word_document(relative_path: str, title: str, paragraphs: list) -> str:
    """
    Buat dokumen Word baru.
    paragraphs: list of string, masing-masing jadi satu paragraf.
    """
    if Document is None:
        return "ERROR: library python-docx belum terinstall. Jalankan: pip install python-docx"

    try:
        target = resolve_safe_path(relative_path)
    except SandboxViolationError as e:
        return f"ERROR: {e}"

    if not str(target).lower().endswith(".docx"):
        target = target.with_suffix(".docx")

    target.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    if title:
        doc.add_heading(title, level=1)
    for para in paragraphs:
        doc.add_paragraph(para)

    doc.save(str(target))
    return f"Dokumen Word '{target.name}' berhasil dibuat di '{relative_path}'."


def read_word_document(relative_path: str) -> str:
    """Baca seluruh isi teks dari dokumen Word."""
    if Document is None:
        return "ERROR: library python-docx belum terinstall. Jalankan: pip install python-docx"

    try:
        target = resolve_safe_path(relative_path)
    except SandboxViolationError as e:
        return f"ERROR: {e}"

    if not target.exists():
        return f"File '{relative_path}' tidak ditemukan."

    doc = Document(str(target))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    return full_text if full_text.strip() else "(Dokumen kosong)"


def create_excel_sheet(relative_path: str, headers: list, rows: list) -> str:
    """
    Buat file Excel baru dengan satu sheet.
    headers: list of string, jadi baris pertama (judul kolom).
    rows: list of list, masing-masing list jadi satu baris data.
    """
    if Workbook is None:
        return "ERROR: library openpyxl belum terinstall. Jalankan: pip install openpyxl"

    try:
        target = resolve_safe_path(relative_path)
    except SandboxViolationError as e:
        return f"ERROR: {e}"

    if not str(target).lower().endswith((".xlsx",)):
        target = target.with_suffix(".xlsx")

    target.parent.mkdir(parents=True, exist_ok=True)

    wb = Workbook()
    ws = wb.active
    if headers:
        ws.append(headers)
    for row in rows:
        ws.append(row)

    wb.save(str(target))
    return f"File Excel '{target.name}' berhasil dibuat di '{relative_path}'."


def read_excel_sheet(relative_path: str) -> str:
    """Baca isi sheet pertama dari file Excel, dikembalikan sebagai teks tabel sederhana."""
    if load_workbook is None:
        return "ERROR: library openpyxl belum terinstall. Jalankan: pip install openpyxl"

    try:
        target = resolve_safe_path(relative_path)
    except SandboxViolationError as e:
        return f"ERROR: {e}"

    if not target.exists():
        return f"File '{relative_path}' tidak ditemukan."

    wb = load_workbook(str(target), data_only=True)
    ws = wb.active

    lines = []
    for row in ws.iter_rows(values_only=True):
        lines.append(" | ".join(str(cell) if cell is not None else "" for cell in row))

    return "\n".join(lines) if lines else "(Sheet kosong)"
